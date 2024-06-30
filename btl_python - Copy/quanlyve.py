from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches
import os
import subprocess
import database
import autocomplete

def show_Ve(frame_noidung):
    for widget in frame_noidung.winfo_children():
        widget.destroy()

    def clear_entries():
        input_ma_kh.delete(0, END)
        input_ma_lichtrinh.delete(0, END)
        input_ma_chongoi.delete(0, END)
        input_trang_thai.delete(0, END)

    def search_ve():
        try:
            keyword = input_timkiem_ve.get()
            rows = database.search_ve(keyword)
            update_table(rows)
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi tìm kiếm: {str(e)}")

    def them_ve(input_ma_kh,input_ma_lichtrinh,input_ma_chongoi,input_trang_thai,addThem):
        try:
            ma_kh = input_ma_kh.get().split(' ')[0]
            ma_lichtrinh = input_ma_lichtrinh.get().split(' ')[0]
            ma_chongoi = input_ma_chongoi.get().split(' ')[0]
            trang_thai = input_trang_thai.get()

            if len(ma_kh) == 0 or len(ma_lichtrinh) == 0 or len(ma_chongoi) == 0 or len(trang_thai) == 0:
                raise ValueError("Vui lòng điền đầy đủ thông tin.")

            # Lấy mã tàu từ mã chỗ ngồi
            ma_tau = database.lay_ma_tau_tu_ma_chongoi(ma_chongoi)
            if ma_tau is None:
                raise ValueError("Mã chỗ ngồi không hợp lệ.")

            # Lấy mã tàu từ mã lịch trình để kiểm tra tính hợp lệ
            ma_tau_lichtrinh = database.lay_ma_tau_tu_ma_lichtrinh(ma_lichtrinh)
            if ma_tau_lichtrinh is None:
                raise ValueError("Mã lịch trình không hợp lệ.")

            # Kiểm tra mã tàu của chỗ ngồi và lịch trình có khớp nhau không
            if ma_tau != ma_tau_lichtrinh:
                raise ValueError("Mã tàu của chỗ ngồi và lịch trình không khớp nhau.")

            # Thêm vé vào cơ sở dữ liệu
            database.them_ve(ma_kh, ma_lichtrinh, ma_chongoi, trang_thai)
            messagebox.showinfo("Thông báo", "Thêm vé thành công")
            clear_entries()
            display_data()
            addThem.destroy()
        except ValueError as ve:
            messagebox.showerror("Lỗi", str(ve))

        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi thêm vé: {str(e)}")

    def fromThem():
        def center_window(window, width, height):
            screen_width = window.winfo_screenwidth()
            screen_height = window.winfo_screenheight()

            x = (screen_width / 2) - (width / 2)
            y = (screen_height / 2) - (height / 2)
            window.geometry('%dx%d+%d+%d' % (width, height, x, y))

        addThem = Toplevel()
        addThem.title('Thêm Tàu')
        addThem.geometry('500x600')
        center_window(addThem, 500, 600)
        addThem.transient(frame_noidung.winfo_toplevel())  # Đảm bảo cửa sổ luôn ở trên cùng của ứng dụng chính
        addThem.attributes("-topmost", True)  # Đảm bảo nó trên tất cả các cửa sổ khác
        addThem.focus_force()  # Đặt trọng tâm vào cửa sổ này

        image = Image.open(r"D:\BAITAPCACMON\BT-CODE\PY\btl_python\img\backgroud.jpg")
        global background_image
        background_image = ImageTk.PhotoImage(image)
        canvas = Canvas(addThem, width=addThem.winfo_screenwidth(), height=addThem.winfo_screenheight())
        canvas.pack(fill=BOTH, expand=YES)

        def resize_image(event):
            new_width = event.width
            new_height = event.height
            resized_image = image.resize((new_width, new_height), Image.LANCZOS)
            global background_image_resized
            background_image_resized = ImageTk.PhotoImage(resized_image)
            canvas.create_image(0, 0, anchor=NW, image=background_image_resized)
            canvas.image = background_image_resized

        canvas.bind("<Configure>", resize_image)
        frame = Frame(addThem, bg="white", width=400, height=330)
        frame.place(relx=0.5, rely=0.5, anchor=CENTER)

        lb = Label(frame, text="Thêm Vé", fg="gray", font=("Arial", 18, "bold"), bg='White')
        lb.place(x=130, y=10)

        lb_ht = Label(frame, text="Mã Khách Hàng", fg="gray", font=("Arial", 10, "bold"), bg='White')
        lb_ht.place(x=30, y=70)

        lb_ns = Label(frame, text="Mã Lịch Trình", fg="gray", font=("Arial", 10, "bold"), bg='White')
        lb_ns.place(x=30, y=120)

        lb_sex = Label(frame, text="Mã Chỗ Ngồi", fg="gray", font=("Arial", 10, "bold"), bg='White')
        lb_sex.place(x=30, y=170)

        lb_sdt = Label(frame, text="Trạng Thái", fg="gray", font=("Arial", 10, "bold"), bg='White')
        lb_sdt.place(x=30, y=220)

        # Nhap
        ma_kh_values = ["{} ({}) ({})".format(row[0], row[1], row[2]) for row in database.get_ma_khach()]
        input_ma_kh = ttk.Combobox(frame,values=ma_kh_values, font=("Arial", 11), width=18)
        input_ma_kh.place(x=150, y=70)

        ma_lichtrinh_values = ["{} ({}) ({}) ({})".format(row[0], row[1], row[2], row[3]) for row in
                               database.get_ma_lt()]
        input_ma_lichtrinh = ttk.Combobox(frame,values=ma_lichtrinh_values, font=("Arial", 11), width=18)
        input_ma_lichtrinh.place(x=150, y=120)

        ma_chongoisua_values = ["{} ({}) ({}) ({}) ({})".format(row[0], row[1], row[2], row[3], row[4]) for row in
                             database.get_ma_cnsua()]
        input_ma_chongoi = ttk.Combobox(frame,values=ma_chongoisua_values,font=("Arial", 11), width=18)
        input_ma_chongoi.place(x=150, y=170)

        input_trang_thai = ttk.Combobox(frame, font=("Arial", 11), state="readonly", values=('DaDat', 'ChuDat'),
                                        width=18)
        input_trang_thai.place(x=150, y=220)

        btn_them = Button(frame, text="Thêm", bg="#57a1f8",command=lambda:them_ve(input_ma_kh,input_ma_lichtrinh,input_ma_chongoi,input_trang_thai,addThem) , fg="White", font=("Arial", 10, "bold"), padx=30, pady=3,
                          borderwidth=0)
        btn_them.place(x=50, y=270)

        btn_dangky = Button(frame, text="Hủy", bg="#57a1f8", fg="White", font=("Arial", 10, "bold"), padx=45, pady=3,
                            borderwidth=0, command=addThem.destroy)
        btn_dangky.place(x=215, y=270)

        addThem.mainloop()

    def sua_ve():
        try:
            selected_row = table_ve.focus()
            data = table_ve.item(selected_row, 'values')

            if not data:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn một vé để sửa.")
                return

            ma_ve = data[0]  # Lấy mã vé từ dữ liệu đã chọn
            ma_kh = input_ma_kh.get().split(' ')[0]
            ma_lichtrinh = input_ma_lichtrinh.get().split(' ')[0]
            ma_chongoi = input_ma_chongoi.get().split(' ')[0]
            trang_thai = input_trang_thai.get()
            ma_tau = database.lay_ma_tau_tu_ma_chongoi(ma_chongoi)
            if len(ma_kh) == 0 or len(ma_lichtrinh) == 0 or len(ma_chongoi) == 0 or len(trang_thai) == 0:
                raise ValueError("Vui lòng điền đầy đủ thông tin.")
            if ma_tau is None:
                raise ValueError("Mã chỗ ngồi không hợp lệ.")

            # Lấy mã tàu từ mã lịch trình để kiểm tra tính hợp lệ
            ma_tau_lichtrinh = database.lay_ma_tau_tu_ma_lichtrinh(ma_lichtrinh)
            if ma_tau_lichtrinh is None:
                raise ValueError("Mã lịch trình không hợp lệ.")

            # Kiểm tra mã tàu của chỗ ngồi và lịch trình có khớp nhau không
            if ma_tau != ma_tau_lichtrinh:
                raise ValueError("Mã tàu của chỗ ngồi và lịch trình không khớp nhau.")
            database.sua_ve(ma_ve, ma_kh, ma_lichtrinh, ma_chongoi, trang_thai)
            messagebox.showinfo("Thông báo", "Cập nhật vé thành công")
            clear_entries()
            display_data()

        except IndexError:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn một vé để sửa.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi cập nhật vé: {str(e)}")

    def xoa_ve():
        try:
            selected_row = table_ve.focus()
            data = table_ve.item(selected_row, 'values')
            ma_ve = data[0]
            database.xoa_ve(ma_ve)
            messagebox.showinfo("Thông báo", "Xóa vé thành công")
            clear_entries()
            display_data()
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi xóa vé: {str(e)}")

    def ChonDuLieuBang(event):
        selected_row = table_ve.focus()
        if not selected_row:
            return

        data = table_ve.item(selected_row, 'values')
        if not data:
            return

        clear_entries()

        # Lấy thông tin từ dòng được chọn
        ma_kh = data[1]
        ma_lichtrinh = data[2]
        ma_chongoi = data[3]
        trang_thai = data[13]

        # Hiển thị mã khách hàng
        ma_kh_values = ["{} ({}) ({})".format(row[0], row[1], row[2]) for row in database.get_ma_khach()]
        input_ma_kh['values'] = ma_kh_values
        for value in ma_kh_values:
            if value.startswith(ma_kh):
                input_ma_kh.set(value)  # Chọn giá trị tương ứng với dòng được chọn

        # Hiển thị mã lịch trình
        ma_lichtrinh_values = ["{} ({}) ({}) ({})".format(row[0], row[1], row[2], row[3]) for row in database.get_ma_lt()]
        input_ma_lichtrinh['values'] = ma_lichtrinh_values
        for value in ma_lichtrinh_values:
            if value.startswith(ma_lichtrinh):
                input_ma_lichtrinh.set(value)  # Chọn giá trị tương ứng với dòng được chọn

        # Hiển thị mã chỗ ngồi
        ma_chongoi_values = ["{} ({}) ({}) ({}) ({})".format(row[0], row[1], row[2], row[3], row[4]) for row in database.get_ma_cnsua()]
        input_ma_chongoi['values'] = ma_chongoi_values
        for value in ma_chongoi_values:
            if value.startswith(ma_chongoi):
                input_ma_chongoi.set(value)  # Chọn giá trị tương ứng với dòng được chọn

        # Hiển thị trạng thái
        input_trang_thai.set(trang_thai)

    def InVe(event):
        selected_row = table_ve.focus()
        data = table_ve.item(selected_row, 'values')
        if not data:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn một vé để in.")
            return

        in_ve_window = Toplevel()
        in_ve_window.title("In vé")
        in_ve_window.geometry("800x400")

        # Đảm bảo đường dẫn đến hình ảnh là chính xác
        image_path = "D:/BAITAPCACMON/BT-CODE/PY/btl_python/img/backgroudIndex.jpg"
        try:
            image = Image.open(image_path)
            background_image = ImageTk.PhotoImage(image)
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể tải hình ảnh: {e}")
            return

        canvas = Canvas(in_ve_window, width=800, height=400)
        canvas.pack(fill="both", expand=True)
        image_width = image.width
        image_height = image.height
        x_center = (800 - image_width) // 2
        y_center = (400 - image_height) // 2
        canvas.create_image(x_center, y_center, anchor=NW, image=background_image)
        in_ve_window.background_image = background_image

        #  cửa sổ nằm ở phía trước
        in_ve_window.transient(frame_noidung)
        in_ve_window.lift()
        in_ve_window.grab_set()

        # Định dạng thông tin vé theo mẫu
        name1 = Label(in_ve_window, text="CÔNG TY CỔ VẬN TẢI", fg="black", font=("Arial", 12, "bold"), bg='White')
        name1.place(x=310, y=50)
        name2 = Label(in_ve_window, text="ĐƯỜNG SÁT", fg="black", font=("Arial", 12, "bold"), bg='White')
        name2.place(x=316, y=70)
        lb_nhom = Label(in_ve_window, text="MR.", bg="white", fg="gray", font=("Arial", 12, "bold"))
        lb_nhom.place(x=420, y=70)
        lb_nhom1 = Label(in_ve_window, text="QK", bg="white", fg="red", font=("Arial", 12, "bold"))
        lb_nhom1.place(x=450, y=70)

        info_text = f"""
        Mã vé/TicketID: {data[0]}
        Ga đi: {data[6]}
        Ga đến: {data[7]}
        Tàu/Train: {data[5]}
        Ngày đi/Date: {data[9]}
        Giờ đi/Time: {data[8]}
        Toa/Coach: {data[3]}
        Chỗ/Seat: {data[10]}
        Loại chỗ/Class: {data[13]}
        Loại vé/Ticket: Người lớn
        Họ tên/Name: {data[4]}
        Giá/Price: {data[12]} VNĐ
        """

        lb_ve_info = Label(in_ve_window, text=info_text, font=("Arial", 12), justify=LEFT, bg="white")
        canvas.create_window(400, 220, window=lb_ve_info, anchor="center")

        def save_to_word():
            try:
                # Tạo tài liệu Word mới
                doc = Document()

                # Thêm hình ảnh vào tài liệu
                doc.add_picture(image_path, width=Inches(5))

                # Thêm thông tin vé vào tài liệu
                doc.add_paragraph("CÔNG TY CỔ VẬN TẢI ĐƯỜNG SẮT\n")
                doc.add_paragraph("MR. QK\n")
                doc.add_paragraph(info_text)

                # Lưu tài liệu
                save_path = "D:/BAITAPCACMON/BT-CODE/PY/btl_python/in_ve.docx"
                doc.save(save_path)
                messagebox.showinfo("Thông báo", "In vé thành công!")

                # Mở tệp Word sau khi lưu
                if os.name == 'nt':  # Windows
                    os.startfile(save_path)
                elif os.name == 'posix':  # macOS or Linux
                    subprocess.call(['open', save_path])
                else:
                    subprocess.call(['xdg-open', save_path])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể lưu vé ra file Word: {e}")

        btn_inve = Button(in_ve_window, command=save_to_word, text="In vé", font=("Arial", 7, "bold"), fg="white",
                          borderwidth=0, bg="#57a1f8", pady=3, padx=29)
        btn_inve.place(y=200)
        canvas.create_window(400, 370, window=btn_inve, anchor="center")
    def display_data():
        try:
            rows = database.Hien_Thi_Ve()
            update_table(rows)
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi hiển thị dữ liệu: {str(e)}")

    def update_table(rows):
        table_ve.delete(*table_ve.get_children())
        for row in rows:
            table_ve.insert("", END, values=row)

    frame_tb = Frame(frame_noidung, bg="White")
    frame_tb.pack(side="top", fill="y")

    frame_chucnang = Frame(frame_noidung, width=300, height=200, bg="White")
    frame_chucnang.pack(side="bottom", fill="both", expand=True)

    frame_timkiem = Frame(frame_tb, bg="White")
    frame_timkiem.pack(side="top", pady=10)

    input_timkiem_ve = Entry(frame_timkiem, font=("Arial", 11), highlightbackground="light blue", highlightthickness=1, borderwidth=0, width=30)
    input_timkiem_ve.grid(row=0, column=0, padx=10, pady=10)

    anh = Image.open(r'D:\BAITAPCACMON\BT-CODE\PY\btl_python\img\icontim.png')
    size = anh.resize((20, 20), Image.LANCZOS)
    global img
    img = ImageTk.PhotoImage(size)
    btn_timkiem_ve = Button(frame_timkiem, image=img, command=search_ve, borderwidth=0)
    btn_timkiem_ve.grid(row=0, column=1, padx=(0, 10))

    columns = ("Mã vé", "Mã khách hàng","Mã lịch trình","Mã Ghế", "Tên khách hàng", "Tên tàu", "Ga khởi hành", "Ga đến", "Thời gian khởi hành", "Ngày Khởi Hành", "Số chỗ ngồi", "Ngày đặt vé", "Giá vé", "Trạng thái")
    table_ve = ttk.Treeview(frame_tb, columns=columns, show="headings")
    for col in columns:
        table_ve.heading(col, text=col)
        table_ve.column(col, minwidth=0, width=76)
    ma_ve = None
    table_ve.pack(fill="both", expand=True)
    for col in columns:
        table_ve.column(col, anchor='center')

    table_ve.bind("<<TreeviewSelect>>", ChonDuLieuBang)
    table_ve.pack(padx=10, pady=10, fill=BOTH, expand=True)
    table_ve.bind("<Double-1>", InVe)

    display_data()

    btn_them = Button(frame_chucnang, command=fromThem, text="Thêm", font=("Arial", 8, "bold"), fg="White", borderwidth=0, bg="#57a1f8", pady=6, padx=45)
    btn_them.place(x=5, y=30)
    btn_sua = Button(frame_chucnang, command=sua_ve, text="Sửa", font=("Arial", 8, "bold"), fg="White", borderwidth=0, bg="#57a1f8", pady=6, padx=45)
    btn_sua.place(x=150, y=30)
    btn_xoa = Button(frame_chucnang, command=xoa_ve, text="Xóa", font=("Arial", 8, "bold"),fg="White", borderwidth=0, bg="#57a1f8", pady=6, padx=45)
    btn_xoa.place(x=290, y=30)

    lb_ma_kh = Label(frame_chucnang, bg="White", text="Mã khách hàng", fg="black", font=("Arial", 10, "bold"))
    lb_ma_kh.place(x=5, y=70)

    ma_kh_values =  ["{} ({}) ({})".format(row[0], row[1], row[2]) for row in database.get_ma_khach()]
    input_ma_kh = ttk.Combobox(frame_chucnang, font=("Arial", 11), values=ma_kh_values, width=18)
    input_ma_kh.place(x=5, y=90)

    lb_ma_lichtrinh = Label(frame_chucnang, bg="White", text="Mã lịch trình", fg="black", font=("Arial", 10, "bold"))
    lb_ma_lichtrinh.place(x=240, y=70)
    ma_lichtrinh_values = ["{} ({}) ({}) ({})".format(row[0], row[1], row[2], row[3]) for row in database.get_ma_lt()]
    input_ma_lichtrinh = ttk.Combobox(frame_chucnang, font=("Arial", 11), values=ma_lichtrinh_values, width=18)
    input_ma_lichtrinh.place(x=240, y=90)

    lb_ma_chongoi = Label(frame_chucnang, bg="White", text="Mã chỗ ngồi", fg="black", font=("Arial", 10, "bold"))
    lb_ma_chongoi.place(x=5, y=130)

    ma_chongoi_values = ["{} ({}) ({}) ({}) ({})".format(row[0], row[1], row[2], row[3], row[4]) for row in database.get_ma_cn()]
    input_ma_chongoi = ttk.Combobox(frame_chucnang, font=("Arial", 11), values=ma_chongoi_values, width=18)
    input_ma_chongoi.place(x=5, y=150)

    lb_trang_thai = Label(frame_chucnang, bg="White", text="Trạng thái", fg="black", font=("Arial", 10, "bold"))
    lb_trang_thai.place(x=240, y=130)
    input_trang_thai = ttk.Combobox(frame_chucnang, font=("Arial", 11), state="readonly", values=('DaDat', 'ChuDat'), width=18)
    input_trang_thai.place(x=240, y=150)
